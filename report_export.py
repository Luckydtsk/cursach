from datetime import date, datetime
from io import BytesIO

from docx import Document
from flask import send_file
from openpyxl import Workbook
from openpyxl.styles import Font


def format_cell(value):
    if value is None:
        return ""
    if isinstance(value, datetime):
        return value.strftime("%d.%m.%Y %H:%M")
    if isinstance(value, date):
        return value.strftime("%d.%m.%Y")
    return str(value)


def build_xlsx(title, headers, rows):
    wb = Workbook()
    ws = wb.active
    ws.title = "Отчет"

    ws.append([title])
    ws["A1"].font = Font(bold=True, size=14)
    ws.append([])

    ws.append(headers)
    for cell in ws[3]:
        cell.font = Font(bold=True)

    for row in rows:
        ws.append([format_cell(v) for v in row])

    for column in ws.columns:
        max_len = 0
        column_letter = column[0].column_letter
        for cell in column:
            if cell.value is not None:
                max_len = max(max_len, len(str(cell.value)))
        ws.column_dimensions[column_letter].width = min(max_len + 2, 50)

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer


def build_docx(title, headers, rows):
    doc = Document()
    doc.add_heading(title, level=1)

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header

    for row in rows:
        row_cells = table.add_row().cells
        for idx, value in enumerate(row):
            row_cells[idx].text = format_cell(value)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def send_report_file(title, headers, rows, filename_stem, file_format):
    file_format = (file_format or "").lower()
    if file_format not in ("xlsx", "docx"):
        raise ValueError("Unsupported format")

    safe_stem = "".join(c if c.isalnum() or c in "-_" else "_" for c in filename_stem)

    if file_format == "xlsx":
        buffer = build_xlsx(title, headers, rows)
        mimetype = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        filename = f"{safe_stem}.xlsx"
    else:
        buffer = build_docx(title, headers, rows)
        mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        filename = f"{safe_stem}.docx"

    return send_file(
        buffer,
        mimetype=mimetype,
        as_attachment=True,
        download_name=filename,
    )
